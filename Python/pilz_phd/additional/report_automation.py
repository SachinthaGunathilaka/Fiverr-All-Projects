import pandas
import pandas as pd
import warnings
from docx import *
from funcoes.createtable import createtable


warnings.simplefilter(action='ignore')

#report_format:
formato_relatorio=pd.read_excel('C:/Users/operador/Instituto PHD/Intranet - Dados e automações (1)/teste/3_report_tables_order.xlsx')

######## CREATES MAIN DATAFRAME

df_responses = pd.DataFrame()
df_questions = pd.DataFrame()
df_analises_final_questoes = pd.DataFrame()


#responses
df_responses = pd.read_excel('C:/Users/operador/Instituto PHD/Intranet - Dados e automações (1)/teste/2_answers.xlsx')


#questions
df_questions= pd.read_excel('C:/Users/operador/Instituto PHD/Intranet - Dados e automações (1)/teste/1_survey_questions.xlsx')

df_responses=df_responses.drop_duplicates()

df_analises=df_responses.merge(df_questions,on=['id_questao','id_resposta'],how='left')
df_analises=df_analises[['id_questionario','id_questao','pergunta','id_resposta','resposta','position','Cod_questao']].sort_values(by=['id_questionario','id_questao'])

df_analises_final=df_analises.groupby(['pergunta','Cod_questao','resposta','id_resposta'],as_index=False)['position'].count().rename(columns={'position':'(n)'})

df_analises_final['%']=100*df_analises_final['(n)']/df_analises_final.groupby('pergunta')['(n)'].transform('sum')
df_analises_final=df_analises_final.sort_values(by=['pergunta','%'])
df_analises_final=df_questions.merge(df_analises_final,on=['id_resposta','pergunta','resposta'],how='left')
df_analises_final=df_analises_final[['pergunta','code','resposta','id_resposta','(n)','%']].rename(columns={'code':'Cod_questao'})
df_analises_final['%']=df_analises_final['%'].fillna(0)
df_analises_final['(n)']=df_analises_final['(n)'].fillna(0).round()


#####creating report

doc = Document()
font = doc.styles['Normal'].font
font.name = 'Calibri (body)'
doc.add_heading('Introdução', level=2)
p = doc.add_paragraph('testando a introdução...')
doc.add_page_break()
df_analises_final_questoes=df_analises_final[['pergunta','Cod_questao']].drop_duplicates()


for reg in formato_relatorio['id']:
    formato_relatorio_for=formato_relatorio.loc[formato_relatorio['id']==reg]
    format = formato_relatorio_for['format'].values[0]

    if format!=str("CROSSED"): #CODE to calculate a BASIC/NPS/SATISFACTION table
        cod_quest=formato_relatorio_for['line'].values[0]
        pergunta=df_analises_final_questoes['pergunta'].loc[df_analises_final_questoes['Cod_questao']==cod_quest].values[0]
        pergunta=pergunta.replace("['","")
        pergunta = pergunta.replace("']", "")
        df_analises_final_for=df_analises_final.loc[df_analises_final['Cod_questao']==cod_quest]
        df_analises_final_for=df_analises_final_for[['resposta','(n)','%']]
        df_analises_final_for['%']=df_analises_final_for['%'].round(1)

        df_analises_final_for.loc['Total']= df_analises_final_for.sum(numeric_only=True, axis=0).round(1)
        df_analises_final_for['%']=df_analises_final_for['%'].astype(str).str.replace('.', ',')
        df_analises_final_for['(n)']=df_analises_final_for['(n)'].astype(str).str.replace('.0', '')
        df_analises_final_for['resposta']=df_analises_final_for['resposta'].fillna('Total')

        doc.add_heading(f'{pergunta}', level=2)
        t = doc.add_table(df_analises_final_for.shape[0] + 1, df_analises_final_for.shape[1])


        ####### CALL FUNCTION WITH DF AND FORMAT
        if format==str("BASIC"):
            createtable(t,df_analises_final_for, 'BASIC')
        elif format==str("NPS"):
            createtable(t, df_analises_final_for, 'NPS')
        elif format==str("SATS"):
            createtable(t, df_analises_final_for, 'SATS')


        doc.add_page_break()

    else:
        cod_quest_linha=formato_relatorio_for['line'].values[0]
        cod_quest_coluna=formato_relatorio_for['column'].values[0]
        pergunta_linha=df_analises_final_questoes['pergunta'].loc[df_analises_final_questoes['Cod_questao']==cod_quest_linha].values[0]
        pergunta_linha=pergunta_linha.replace("['","")
        pergunta_linha = pergunta_linha.replace("']", "")
        pergunta_coluna=df_analises_final_questoes['pergunta'].loc[df_analises_final_questoes['Cod_questao']==cod_quest_coluna].values[0]
        pergunta_coluna=pergunta_coluna.replace("['","")
        pergunta_coluna = pergunta_coluna.replace("']", "")
        pergunta = pergunta_linha+str(' x ')+pergunta_coluna

        df_analises_final_cruz_1=df_analises.loc[(df_analises['Cod_questao'] == cod_quest_linha)]
        df_analises_final_cruz_1['tipo']=str('linha')
        df_analises_final_cruz_2=df_analises.loc[(df_analises['Cod_questao'] == cod_quest_coluna)]
        df_analises_final_cruz_2['tipo']=str('coluna')
        df_analises_final_cruz=df_analises_final_cruz_1.append(df_analises_final_cruz_2)

        df_analises_final_cruz=df_analises_final_cruz.pivot(index=['id_questionario'],columns='tipo',values='resposta').reset_index()

        df_cruzamento=pandas.pivot_table(df_analises_final_cruz,index='linha',columns='coluna', margins=True, margins_name='Global (%)',aggfunc='count')
        df_cruzamento = df_cruzamento.loc[df_cruzamento.index!='Global (%)']
        df_cruzamento = (100. * df_cruzamento / df_cruzamento.sum()).round(1)

        df_cruzamento=df_cruzamento.reset_index()
        df_cruzamento.columns = df_cruzamento.columns.droplevel(0)
        df_cruzamento=df_cruzamento.rename(columns={'':'linha'})

        df_cruzamento_final=df_analises_final.loc[df_analises_final['Cod_questao']==cod_quest_linha]
        df_cruzamento_final=df_cruzamento_final[['resposta']].rename(columns={'resposta':'linha'})
        df_cruzamento_final=df_cruzamento_final.merge(df_cruzamento,on=['linha'],how='left').fillna(0)

        df_cruzamento_final.loc['Total']= df_cruzamento_final.sum(numeric_only=True, axis=0).round(1)

        df_cruzamento_final['linha']=df_cruzamento_final['linha'].fillna('Total')



        ####### HERE I CREATE THE TABLE
        doc.add_heading(f'{pergunta}', level=3)
        t = doc.add_table(df_cruzamento_final.shape[0] + 1, df_cruzamento_final.shape[1])

        createtable(t,df_cruzamento_final, 'CROSSED')


        doc.add_page_break()


doc.save(r'C:/Users/operador/Instituto PHD/Intranet - Dados e automações (1)/teste/demo.docx')
