from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
import pandas
import pandas as pd
import warnings
from docx import *
from docx.enum.text import WD_LINE_SPACING
import docx
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx.shared import Pt, Inches

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum import text as textAlign

def createTable(word_doc, word_doc_name, table_type, data):

    def make_rows_bold(R, G, B, *rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(R, G, B)
                        run.font.bold = True


    def make_cols_bold(R, G, B, table, col_index):
        for row in table.rows[1:]:
            for paragraph in row.cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(R, G, B)
                    run.font.bold = True

    temp_doc_name = "generated_doc.docx"
    basic_template_doc = "basic_template.docx"
    nps_template_doc = "nps_template.docx"
    sats_template_doc = "sats_template.docx"

    if table_type != "CROSSED":
        word_doc.save(word_doc_name)

    doc = None
    if table_type == "BASIC":
        print("BASIC")
        doc = DocxTemplate(basic_template_doc)
        context = {
            'title_0': data.columns[0],
            'title_1': data.columns[1],
            'title_2': data.columns[2],

            'data_0_0': data.iloc[0][0],
            'data_0_1': data.iloc[0][1],
            'data_0_2': data.iloc[0][2],
        }

        doc.render(context)
        doc.save(temp_doc_name)

        master = Document_compose(word_doc_name)
        composer = Composer(master)
        #filename_second_docx is the name of the second docx file
        doc2 = Document_compose(temp_doc_name)
        #append the doc2 into the master using composer.append function
        composer.append(doc2)
        #Save the combined docx with a name
        composer.save(word_doc_name)
        doc = Document(word_doc_name)

        table = doc.tables[-1]

        for i in range(1, data.shape[0]):
            row = table.add_row()
            # row.height = Pt(1)
            cells = row.cells

            cells[0].text = str(data.iloc[i][0])
            cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            cells[1].text = str(data.iloc[i][1])
            cells[1].paragraphs[0].alignment = textAlign.WD_ALIGN_PARAGRAPH.CENTER
            cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            cells[2].text = str(data.iloc[i][2])
            cells[2].paragraphs[0].alignment = textAlign.WD_ALIGN_PARAGRAPH.CENTER
            cells[2].paragraphs[0].paragraph_format.space_after = Pt(0)
            cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for row in table.rows:
            row.height = Inches(0.25)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        make_rows_bold(0, 0, 0, table.rows[data.shape[0]])

    elif table_type == "NPS":

        print("NPS")
        doc = DocxTemplate(nps_template_doc)
        percentages = data.iloc[:, 2].copy().map(lambda x: float(x.replace(",", ".")))

        q1 = percentages[:7].sum().round(2)
        q2 = percentages[7:9].sum().round(2)
        q3 = percentages[9:11].sum().round(2)
        tq = q1 + q2 + q3

        nps = q3 - q1

        context = {
            'd0': data.iloc[0, 2],
            'd1': data.iloc[1, 2],
            'd2': data.iloc[2, 2],
            'd3': data.iloc[3, 2],
            'd4': data.iloc[4, 2],
            'd5': data.iloc[5, 2],
            'd6': data.iloc[6, 2],
            'd7': data.iloc[7, 2],
            'd8': data.iloc[8, 2],
            'd9': data.iloc[9, 2],
            'd10': data.iloc[10, 2],
            'nd': data.iloc[11, 2].replace(".", ","),
            'td': data.iloc[12, 2].replace(".", ","),

            'q1': str(q1).replace(".", ","),
            'q2': str(q2).replace(".", ","),
            'q3': str(q3).replace(".", ","),
            'nq': data.iloc[11, 2].replace(".", ","),
            'tq': str(tq).replace(".", ","),
            'nps': str(nps).replace(".", ",")
        }

        doc.render(context)
        doc.save(temp_doc_name)

        master = Document_compose(word_doc_name)
        composer = Composer(master)
        #filename_second_docx is the name of the second docx file
        doc2 = Document_compose(temp_doc_name)
        #append the doc2 into the master using composer.append function
        composer.append(doc2)
        #Save the combined docx with a name
        composer.save(word_doc_name)
        doc = Document(word_doc_name)


    elif table_type == "SATS":
        print("SATS")
        doc = DocxTemplate(sats_template_doc)

        percentages = data.iloc[:, 2].copy().map(lambda x: float(x.replace(",", ".")))
        q1 = percentages[:2].sum().round(2)
        q2 = percentages[3:5].sum().round(2)

        fav = (q2 * 100 / percentages[:5].sum()).round(2)

        context = {
            'title_1': data.columns[0],
            'title_2': data.columns[2],

            'd0': data.iloc[0, 0],
            'd1': data.iloc[1, 0],
            'd2': data.iloc[2, 0],
            'd3': data.iloc[3, 0],
            'd4': data.iloc[4, 0],
            'd5': data.iloc[5, 0],
            'd6': data.iloc[6, 0],

            'v0': data.iloc[0, 2],
            'v1': data.iloc[1, 2],
            'v2': data.iloc[2, 2],
            'v3': data.iloc[3, 2],
            'v4': data.iloc[4, 2],
            'v5': data.iloc[5, 2],
            'v6': data.iloc[6, 2],

            'q1': str(q1).replace(".", ","),
            'q2': str(q2).replace(".", ","),
            'fav': str(fav).replace(".", ","),
        }

        doc.render(context)
        doc.save(temp_doc_name)

        master = Document_compose(word_doc_name)
        composer = Composer(master)
        #filename_second_docx is the name of the second docx file
        doc2 = Document_compose(temp_doc_name)
        #append the doc2 into the master using composer.append function
        composer.append(doc2)
        #Save the combined docx with a name
        composer.save(word_doc_name)
        doc = Document(word_doc_name)

    elif table_type == "CROSSED":
        print("CROSSED")

        doc = word_doc
        num_rows = len(data.axes[0]) + 1
        num_cols = len(data.axes[1])
        # Add a table to the document
        table = doc.add_table(rows=num_rows, cols=num_cols)

        col_id = 0
        row_id = -1
        for row in table.rows:
            col_id = 0
            for cell in row.cells:
                tc = cell._tc
                tc_pr = tc.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'double')
                top.set(qn('w:sz'), '12')
                top.set(qn('w:space'), '0')
                tc_borders.append(top)
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'double')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:space'), '0')
                tc_borders.append(bottom)
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'double')
                left.set(qn('w:sz'), '12')
                left.set(qn('w:space'), '0')
                tc_borders.append(left)
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'double')
                right.set(qn('w:sz'), '12')
                right.set(qn('w:space'), '0')
                tc_borders.append(right)
                tc_pr.append(tc_borders)

                value = ""
                if row_id == -1:
                    value = str(data.columns[col_id]).replace(".", ",")
                else:
                    value = str(data.iloc[row_id, col_id]).replace(".", ",")

                cell.text = str(value)
                cell.paragraphs[0].alignment = textAlign.WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                col_id += 1

            row_id += 1

        # dark_green = "0c5c2b"
        # light_green = "c6d5a1"


        for i in range(num_cols):
            background_style = parse_xml(r'<w:shd {} w:fill="0c5c2b"/>'.format(nsdecls('w')))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(background_style)

        for i in range(1, num_rows):
            background_style = parse_xml(r'<w:shd {} w:fill="c6d5a1"/>'.format(nsdecls('w')))
            table.rows[i].cells[num_cols - 1]._tc.get_or_add_tcPr().append(background_style)

        cells = table.rows[0].cells

        cells[0].text = "Column→\nLine ↓"
        cells[0].paragraphs[0].alignment = textAlign.WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        make_rows_bold(255, 255, 255, table.rows[0])
        make_rows_bold(0, 0, 0, table.rows[num_rows - 1])
        make_cols_bold(0, 0, 0, table, 0)
        make_cols_bold(0, 0, 0, table, num_cols - 1)


    return doc


final_output_filename = "demo.docx"


warnings.simplefilter(action='ignore')

#report_format:
formato_relatorio = pd.read_excel('3_report_tables_order.xlsx')
######## CREATES MAIN DATAFRAME

df_responses = pd.DataFrame()
df_questions = pd.DataFrame()
df_analises_final_questoes = pd.DataFrame()

#responses
df_responses = pd.read_excel('2_answers.xlsx')

#questions
df_questions = pd.read_excel('1_survey_questions.xlsx')

df_responses = df_responses.drop_duplicates()

df_analises = df_responses.merge(df_questions, on=['id_questao', 'id_resposta'], how='left')
df_analises = df_analises[
    ['id_questionario', 'id_questao', 'pergunta', 'id_resposta', 'resposta', 'position', 'Cod_questao']].sort_values(
    by=['id_questionario', 'id_questao'])

df_analises_final = df_analises.groupby(['pergunta', 'Cod_questao', 'resposta', 'id_resposta'], as_index=False)[
    'position'].count().rename(columns={'position': '(n)'})

df_analises_final['%'] = 100 * df_analises_final['(n)'] / df_analises_final.groupby('pergunta')['(n)'].transform('sum')
df_analises_final = df_analises_final.sort_values(by=['pergunta', '%'])
df_analises_final = df_questions.merge(df_analises_final, on=['id_resposta', 'pergunta', 'resposta'], how='left')
df_analises_final = df_analises_final[['pergunta', 'code', 'resposta', 'id_resposta', '(n)', '%']].rename(
    columns={'code': 'Cod_questao'})
df_analises_final['%'] = df_analises_final['%'].fillna(0)
df_analises_final['(n)'] = df_analises_final['(n)'].fillna(0).round()

#####creating report

doc = Document()
font = doc.styles['Normal'].font
font.name = 'Calibri (body)'
doc.add_heading('Introdução', level=2)
p = doc.add_paragraph('testando a introdução...')
doc.add_page_break()
df_analises_final_questoes = df_analises_final[['pergunta', 'Cod_questao']].drop_duplicates()


for reg in formato_relatorio['id']:
    formato_relatorio_for = formato_relatorio.loc[formato_relatorio['id'] == reg]
    format = formato_relatorio_for['format'].values[0]

    if format != str("CROSSED"):  #CODE to calculate a BASIC/NPS/SATISFACTION table
        cod_quest=formato_relatorio_for['line'].values[0]
        pergunta=df_analises_final_questoes['pergunta'].loc[df_analises_final_questoes['Cod_questao']==cod_quest].values[0]
        pergunta=pergunta.replace("['","")
        pergunta = pergunta.replace("']", "")
        df_analises_final_for=df_analises_final.loc[df_analises_final['Cod_questao']==cod_quest]
        df_analises_final_for=df_analises_final_for[['resposta','(n)','%']]
        df_analises_final_for['%']=df_analises_final_for['%'].round(1)

        df_analises_final_for.loc['Total']= df_analises_final_for.sum(numeric_only=True, axis=0).round(1)
        df_analises_final_for['%']=df_analises_final_for['%'].astype(str).str.replace('.', ',')
        df_analises_final_for['(n)']=df_analises_final_for['(n)'].astype(str).str.replace('.0', '')
        df_analises_final_for['resposta']=df_analises_final_for['resposta'].fillna('Total')

        doc.add_heading(f'{pergunta}', level=2)

        if format == str("BASIC"):
            doc = createTable(doc, final_output_filename, "BASIC", df_analises_final_for)

        elif format == str("NPS"):
            doc = createTable(doc, final_output_filename, "NPS", df_analises_final_for)

        elif format == str("SATS"):
            doc = createTable(doc, final_output_filename, "SATS", df_analises_final_for)

        doc.add_page_break()

    else:
        cod_quest_linha=formato_relatorio_for['line'].values[0]
        cod_quest_coluna=formato_relatorio_for['column'].values[0]
        pergunta_linha=df_analises_final_questoes['pergunta'].loc[df_analises_final_questoes['Cod_questao']==cod_quest_linha].values[0]
        pergunta_linha=pergunta_linha.replace("['","")
        pergunta_linha = pergunta_linha.replace("']", "")
        pergunta_coluna=df_analises_final_questoes['pergunta'].loc[df_analises_final_questoes['Cod_questao']==cod_quest_coluna].values[0]
        pergunta_coluna=pergunta_coluna.replace("['","")
        pergunta_coluna = pergunta_coluna.replace("']", "")
        pergunta = pergunta_linha+str(' x ')+pergunta_coluna

        df_analises_final_cruz_1=df_analises.loc[(df_analises['Cod_questao'] == cod_quest_linha)]
        df_analises_final_cruz_1['tipo']=str('linha')
        df_analises_final_cruz_2=df_analises.loc[(df_analises['Cod_questao'] == cod_quest_coluna)]
        df_analises_final_cruz_2['tipo']=str('coluna')
        df_analises_final_cruz=df_analises_final_cruz_1.append(df_analises_final_cruz_2)

        df_analises_final_cruz=df_analises_final_cruz.pivot(index=['id_questionario'],columns='tipo',values='resposta').reset_index()

        df_cruzamento=pandas.pivot_table(df_analises_final_cruz,index='linha',columns='coluna', margins=True, margins_name='Global (%)',aggfunc='count')
        df_cruzamento = df_cruzamento.loc[df_cruzamento.index!='Global (%)']
        df_cruzamento = (100. * df_cruzamento / df_cruzamento.sum()).round(1)

        df_cruzamento=df_cruzamento.reset_index()
        df_cruzamento.columns = df_cruzamento.columns.droplevel(0)
        df_cruzamento=df_cruzamento.rename(columns={'':'linha'})

        df_cruzamento_final=df_analises_final.loc[df_analises_final['Cod_questao']==cod_quest_linha]
        df_cruzamento_final=df_cruzamento_final[['resposta']].rename(columns={'resposta':'linha'})
        df_cruzamento_final=df_cruzamento_final.merge(df_cruzamento,on=['linha'],how='left').fillna(0)

        df_cruzamento_final.loc['Total']= df_cruzamento_final.sum(numeric_only=True, axis=0).round(1)

        df_cruzamento_final['linha']=df_cruzamento_final['linha'].fillna('Total')


        doc.add_heading(f'{pergunta}', level=3)
        doc = createTable(doc, final_output_filename, "CROSSED", df_cruzamento_final)
        doc.add_page_break()

doc.save(final_output_filename)