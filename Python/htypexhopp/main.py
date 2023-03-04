import docx
import re

filename = "doc.docx"
doc = docx.Document(filename)

out = doc

# for i in range(len(out.paragraphs)):
#     out.paragraphs[i] = None

allTables = doc.tables

for activeTable in allTables:
    activeTable._element.getparent().remove(activeTable._element)

allTables = doc.tables

out.save("a.docx")


#
class Headers:
    def __init__(self):
        self.headers = []

    def addHeader(self, number, text):
        self.headers.append([number, text])

    def getLevels(self, header_num):
        return len(header_num.split(".")) - 2

    def display(self):
        for [header_num, text] in self.headers:
            print("\t" * self.getLevels(header_num) + header_num, text)


def iter_heading(paragraphs):
    for paragraph in paragraphs:
        isItHeading = re.match('Heading ([1-9])', paragraph.style.name)
        if isItHeading:
            yield int(isItHeading.groups()[0]), paragraph


def addHeaderNumbering(document):
    hNums = [-1, -1, -1, -1, -1, -1, -1, -1, -1]
    for index, hx in iter_heading(document.paragraphs):
        if hx.text.strip() == "" and index > 1:
            continue

        for i in range(index + 1, 9):
            if i > 2:
                hNums[i] = 0
            else:
                hNums[i] = -1

        hNums[index] += 1

        header_num = ""

        if index < 4:
            for i in range(1, index + 1):
                if i == 1:
                    header_num += chr(64 + hNums[i]) + "."
                else:
                    header_num += "%d." % hNums[i]

            headers.addHeader(header_num, hx.text.strip())
            # hx.text = header_num + " " + hx.text.strip()
            # headers.append(hx.text.strip())
            # print(hx.text)


headers = Headers()
addHeaderNumbering(doc)
headers.display()
#
#
# def get_para_data(output_doc_name, paragraph):
#     """
#     Write the run to the new file and then set its font, bold, alignment, color etc. data.
#     """
#
#     output_para = output_doc_name.add_paragraph()
#     for run in paragraph.runs:
#         output_run = output_para.add_run(run.text)
#         # Run's bold data
#         output_run.bold = run.bold
#         # Run's italic data
#         output_run.italic = run.italic
#         # Run's underline data
#         output_run.underline = run.underline
#         # Run's color data
#         output_run.font.color.rgb = run.font.color.rgb
#         # Run's font data
#         output_run.style.name = run.style.name
#     # Paragraph's alignment data
#     output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
#
#
# output_doc = docx.Document()
#
# for para in doc.paragraphs:
#     get_para_data(output_doc, para)
# output_doc.save('OutputDoc.docx')


p_s = soup.findAll("w:p")
is_find = False
for p in p_s:
    if is_find:
        break

    ppr_s = p.findAll("w:pPr")
    if len(ppr_s) == 0:
        continue

    for ppr in ppr_s:
        p_style = ppr.findAll("w:pStyle", {"w:val": "Heading3"})[13]

        # if len(p_style) == 0:
        #     continue

        is_find = True
        header_content = str(p_style[0].parent.parent)
        pos = header_content.find("paraId")

        end_pos = pos + 8
        while True:
            if header_content[end_pos] == '\"':
                break

            end_pos += 1

        print(header_content[pos + 8: end_pos])

        break
