import os
from tkinter.ttk import Style

if __name__ == '__main__':

    import docx
    import re
    from tkinter import *
    from tkinter import ttk
    from bs4 import BeautifulSoup
    import shutil
    import zipfile
    from tkinter import filedialog
    from tkinter.filedialog import asksaveasfile

    doc_file_path = "input_doc.docx"
    xml_file_path = "Word Document/word/document.xml"
    zip_file_path = "input_zip.zip"
    output_zip_path = "output.zip"
    output_doc_path = "output.docx"

    header_vars = []
    chapter_names = []
    chapter_names_vars = []
    order_vars = []
    order_error = False
    error = None
    order_list = []
    chapter_indices = []
    empty_chaps_vars = []
    empty_chaps_order_vars = []


    def clearFiles():
        for file in [zip_file_path, output_zip_path, output_doc_path, doc_file_path]:
            if os.path.exists(file):
                os.remove(file)

        try:
            shutil.rmtree("Word Document")

        except:
            pass


    clearFiles()


    def initDocFile(filename):
        shutil.copyfile(filename, zip_file_path)

        with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
            zip_ref.extractall("Word Document")


    def CreateDoc():
        shutil.make_archive("output", 'zip', "Word Document")
        os.rename('output.zip', 'output.docx')

        f = asksaveasfile(initialfile='Untitled.docx',
                          defaultextension=".docx", filetypes=[("Doc Files", "*.docx")])

        save_file_name = f.name
        f.close()

        shutil.copy("output.docx", save_file_name)


    # initDocFile(doc_file_path)

    # xml_file = open(xml_file_path, 'r')
    # content = xml_file.read()
    # soup = BeautifulSoup(content, 'xml')
    # xml_file.close()

    class Headers:
        def __init__(self):
            self.headers = []

        def addHeader(self, number, text):
            self.headers.append([number, text])

        def getLevels(self, header_num):
            return len(header_num.split(".")) - 2

        def display(self):
            for [header_num, text] in self.headers:
                print("\t" * self.getLevels(header_num) + header_num, text)

        def getChapters(self):
            chapters = []
            for i in range(len(headers.headers)):
                if self.getLevels(self.headers[i][0]) == 0:
                    chapters.append(self.headers[i])
            return chapters

        def getHeaderTextUI(self, header):
            return "\t" * self.getLevels(header[0]) + header[0] + header[1]

        def getHeaderPath(self, index):
            path = []

            header_nums = self.headers[index][0].split(".")

            for i in range(len(header_nums) - 1):
                if i == 0:
                    path.append(ord(header_nums[i]) - 64)
                elif i == 1:
                    path.append(int(header_nums[i]))
                else:
                    path.append(int(header_nums[i]) - 1)

            return path


    headers = Headers()


    def iter_heading(paragraphs):
        for paragraph in paragraphs:
            isItHeading = re.match('Heading ([1-9])', paragraph.style.name)
            if isItHeading:
                yield int(isItHeading.groups()[0]), paragraph


    def addHeaderNumbering(document):
        hNums = [-1, -1, -1, -1, -1, -1, -1, -1, -1]
        for index, hx in iter_heading(document.paragraphs):

            for i in range(index + 1, 9):
                if i > 2:
                    hNums[i] = 0
                else:
                    hNums[i] = -1

            hNums[index] += 1

            header_num = ""

            if index < 4:
                for i in range(1, index + 1):
                    if i == 1:
                        header_num += chr(64 + hNums[i]) + "."
                    else:
                        header_num += "%d." % hNums[i]

                headers.addHeader(header_num, hx.text.strip())

        # return headers


    def getHeaderNum(header_details):
        iter_level = 0
        count = -1
        selected = -1
        for i in range(len(headers.headers)):
            level = headers.getLevels(headers.headers[i][0])

            if iter_level == level:
                count += 1
                if header_details[iter_level] == count:
                    iter_level += 1
                    count = -1

                    if iter_level == len(header_details):
                        selected = i
                        break

        return selected


    def getParaIds():
        xml_file = open(xml_file_path, 'r')
        content = xml_file.read()
        soup = BeautifulSoup(content, 'xml')
        xml_file.close()

        header_para_ids = []
        p_s = soup.findAll("w:p")
        index = 0

        for p in p_s:
            # print("p")
            ppr_s = p.findAll("w:pPr")
            if len(ppr_s) == 0:
                continue

            for ppr in ppr_s:
                # print("ppr")
                p_style = ppr.findAll("w:pStyle",
                                      {"w:val": re.compile("Heading1|Titre1|Heading2|Titre2|Heading3|Titre3")})
                if len(p_style) == 0:
                    continue

                index += 1

                # print("k")
                header_content = str(p_style[0].parent.parent)
                # print("header", header_content)
                pos = header_content.find("paraId")

                end_pos = pos + 8
                while True:
                    if header_content[end_pos] == '\"':
                        break

                    end_pos += 1

                header_para_ids.append(header_content[pos + 8: end_pos])
                # print(header_content[pos + 8:end_pos])

        return header_para_ids


    # doc = docx.Document(doc_file_path)
    # headers = Headers()
    # addHeaderNumbering(doc)



    def generateXML():
        global chapter_indices
        global header_vars

        try:
            shutil.rmtree("Word Document")

        except:
            pass

        with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
            zip_ref.extractall("Word Document")

        for file in [output_doc_path, output_zip_path]:
            if os.path.exists(file):
                os.remove(file)

        if os.path.exists(xml_file_path):
            ###
            ###

            selectedData = []
            for i in range(len(header_vars)):
                if header_vars[i].get() == 1:
                    selectedData.append([i, headers.getHeaderPath(i)])

            print("selected : ", selectedData)

            ###
            ###

            input_data = []
            sortList()
            print("chapter indices :", chapter_indices)

            for index in chapter_indices:
                for i in range(len(selectedData)):
                    if index == selectedData[i][1][0]:
                        input_data.append(selectedData[i])

            print("input data : ", input_data)
            header_para_ids = getParaIds()
            # print(header_para_ids)

            xml_string = open(xml_file_path).read()
            body_start = re.search("<w:body[\s\S]*>", xml_string).start()

            body_end = body_start + 5
            while True:
                if xml_string[body_end] == '<':
                    break
                body_end += 1

            out_content = xml_string[:body_end]

            # out_content += open("empty_template.txt").read()
            for input_header_details in input_data:
                # print(getHeaderNum(input_header_details[1]))
                # print(len(header_para_ids))
                start_paraId = header_para_ids[getHeaderNum(input_header_details[1])]

                end_paraId = None
                if input_header_details[0] != len(headers.headers) - 1:
                    end_paraId = header_para_ids[getHeaderNum(input_header_details[1]) + 1]

                    # print(start_paraId, end_paraId)

                header_start = re.search(f"paraId=\"{start_paraId}\"", xml_string).start()

                header_1_start = header_start
                while True:
                    if xml_string[header_1_start] == "<":
                        break

                    header_1_start -= 1

                if input_header_details[0] != len(headers.headers) - 1:
                    header_end = re.search(f"paraId=\"{end_paraId}\"", xml_string).start()
                    header_2_start = header_end

                    while True:
                        if xml_string[header_2_start] == "<":
                            break

                        header_2_start -= 1

                    out_content += xml_string[header_1_start:header_2_start]

                else:

                    last_pos = re.search("</w:body[\s\S]*>", xml_string).start()
                    out_content += xml_string[header_1_start:last_pos]

            end_content = re.search("</w:body[\s\S]*>", xml_string).start()
            out_content += xml_string[end_content:]

            out_xml = open(xml_file_path, 'w')

            out_xml.write(out_content)
            out_xml.close()

            # sortList()

            CreateDoc()


    select_button_status = 0


    def selectAll():
        global select_button_status

        if select_button_status == 0:
            select_button_status = 1
            for i in range(len(header_vars)):
                header_vars[i].set(1)

        else:
            select_button_status = 0
            for i in range(len(header_vars)):
                header_vars[i].set(0)


    def browseFiles():
        global header_vars
        global headers
        global doc_file_path

        clearFiles()
        for widget in scrollable_frame.winfo_children():
            widget.destroy()

        header_vars = []
        headers = Headers()

        filename = filedialog.askopenfilename(initialdir="/",
                                              title="Select a File",
                                              filetypes=(("Doc files",
                                                          "*.docx"),
                                                         ))

        # Change label contents
        label_file_explorer.configure(text="File Opened: " + filename)
        shutil.copy(filename, doc_file_path)

        Button(scrollable_frame, text="Select/Deselect All", command=selectAll).pack(anchor="w", pady=20, padx=100)

        initDocFile(doc_file_path)

        doc_file_path = "./input_doc.docx"
        doc = docx.Document(doc_file_path)
        # headers = Headers()
        addHeaderNumbering(doc)

        for x in range(len(headers.headers)):
            level = headers.getLevels(headers.headers[x][0])
            header_vars.append(IntVar())
            header_text = headers.headers[x][0] + " " + headers.headers[x][1]

            Checkbutton(scrollable_frame, text=header_text, variable=header_vars[x]).pack(anchor="w",
                                                                                          padx=level * 40 + 700)

            # Checkbutton(frame_r, text=header_text, variable=header_vars[x]).pack(anchor="w",
            #                                                                               padx=level * 40 + 100)
            # placement += 20
            # row_num += 1


    def sortList():
        swapped = False
        for i in range(len(order_list) - 1):
            for j in range(0, len(order_list) - i - 1):
                if order_list[j] > order_list[j + 1]:
                    swapped = True
                    order_list[j], order_list[j + 1] = order_list[j + 1], order_list[j]
                    chapter_indices[j], chapter_indices[j + 1] = chapter_indices[j + 1], chapter_indices[j]

            if not swapped:
                return

        # print(order_list)
        # print(chapter_indices)


    def changeChapterNames():
        global chapter_names_vars
        global order_error
        global order_vars
        global error
        global order_list

        try:
            order_error = False
            if error:
                error.destroy()

            order_list = []
            for order_var in order_vars:
                order_list.append(order_var.get())

            print(order_list)
            for i in range(len(order_list)):
                if order_list.count(order_list[i]) > 1:
                    order_error = True
                    break

            if order_error:
                error = Label(frame_r, text=f"Invalid ordering of the chapters", relief="solid", pady=5,
                              font=("Courier", 10),
                              bg="red")
                error.grid(row=len(chapter_names_vars) + 2, column=1, pady=10)

        except:
            error = Label(frame_r, text=f"Invalid ordering of the chapters", relief="solid", pady=5,
                          font=("Courier", 10),
                          bg="red")

            error.grid(row=len(chapter_names_vars) + 2, column=1, pady=10)


    def createChapterNames():
        global chapter_names
        global chapter_names_vars
        global order_vars
        global chapter_indices

        chapter_names = []
        chapter_indices = []

        for i in range(len(header_vars)):
            # print(headers.headers[i])
            if header_vars[i].get() == 1:
                if len(headers.getHeaderPath(i)) == 1:
                    chapter_names.append(headers.headers[i])
                    chapter_indices.append(headers.getHeaderPath(i)[0])

        # print(chapter_names)

        order_vars = []

        for _ in chapter_names:
            chapter_names_vars.append(StringVar())

        for i in range(len(chapter_names)):
            order_vars.append(IntVar())
            order_vars[i].set(i + 1)

            Label(frame_r, text=f"Chapter {chapter_names[i][0]}", borderwidth=2, relief="solid", pady=5,
                  font=("Courier", 10),
                  bg="orange").grid(row=i + 1, column=0, pady=10)

            Entry(frame_r, textvariable=chapter_names_vars[i], font=('calibre', 10, 'normal'), width=50).grid(row=i + 1,
                                                                                                              column=1,
                                                                                                              pady=10,
                                                                                                              padx=10)

            Entry(frame_r, textvariable=order_vars[i], font=('calibre', 10, 'normal'), width=5).grid(row=i + 1,
                                                                                                     column=2,
                                                                                                     pady=10)
            chapter_names_vars[i].set(chapter_names[i][1])

        def addNewChapter():
            global empty_chaps_vars
            global empty_chaps_order_vars

            empty_chaps_vars.append(StringVar())
            empty_chaps_order_vars.append(IntVar())

            position = len(order_vars) + len(empty_chaps_vars)
            # order_vars.append(IntVar())
            # chapter_names_vars.append(StringVar())
            # order_vars[-1].set()

            Label(frame_r, text=f"New Chapter", borderwidth=2, relief="solid", pady=5,
                  font=("Courier", 10),
                  bg="orange").grid(row=position + 1, column=0, pady=10)

            Entry(frame_r, textvariable=empty_chaps_vars[-1], font=('calibre', 10, 'normal'), width=50).grid(
                row=position + 1,
                column=1,
                pady=10,
                padx=10)

            Entry(frame_r, textvariable=empty_chaps_order_vars[-1], font=('calibre', 10, 'normal'), width=5).grid(
                row=position + 1,
                column=2,
                pady=10)

            empty_chaps_vars[-1].set("Undefined")

        submit_btn = Button(frame_r, text='Change Chapter Names', command=changeChapterNames, bg="green", fg="white",
                            font=('calibre', 12, 'bold'), padx=10, pady=5).grid(row=len(chapter_names) + 1, column=0,
                                                                                pady=30)

        add_new = Button(frame_r, text='Add New Chapter', command=addNewChapter, bg="blue", fg="white",
                         font=('calibre', 12, 'bold'), padx=10, pady=5).grid(row=len(chapter_names) + 1, column=1,
                                                                             pady=30)
        # btn2.grid(row=1, column=3, pady=10, padx=100)


    def showChapters():

        # if len(headers.getHeaderPath(i)) == 1 and headers.headers:
        #
        #         print(headers.headers[i])

        for widget in frame_r.winfo_children():
            widget.destroy()

        refresh_button = Button(frame_r, text='Show Selected Chapter Names', command=showChapters, bg="green",
                                fg="white",
                                font=('calibre', 12, 'bold'), padx=10, pady=5).grid(row=0, column=0,
                                                                                    pady=30)
        createChapterNames()


    root = Tk()
    root.title('Word Document Scraper')
    width = root.winfo_screenwidth()
    height = root.winfo_screenheight()
    root.geometry("%dx%d" % (width, height))

    container = ttk.Frame(root)
    right = ttk.Frame(root)

    canvas = Canvas(container, width=width - 100, height=height - 100)

    # left = ttk.Frame(container)
    scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
    scrollable_frame = ttk.Frame(canvas)
    frame_r = ttk.Frame(canvas)

    # left.pack(side="left", fill="both", expand=True)
    # right.pack(side="right", fill="both", expand=True)

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )

    frame_r.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.create_window((0, 0), window=frame_r, anchor="nw")

    refresh_button = Button(frame_r, text='Show Selected Chapter Names', command=showChapters, bg="green", fg="white",
                            font=('calibre', 12, 'bold'), padx=10, pady=5).grid(row=0, column=0,
                                                                                pady=30)

    canvas.configure(yscrollcommand=scrollbar.set)

    label_file_explorer = Label(root,
                                text="Select Doc File",
                                width=100, height=4,
                                fg="blue")

    # button_explore = Button(root,
    #                         text="Browse Files",
    #                         command=browseFiles)

    button_explore = Button(root, text='Browse Files', command=browseFiles, bg="green", fg="white",
                            font=('calibre', 12, 'bold'), padx=10, pady=5)

    # remove later
    # browseFiles()

    label_file_explorer.pack()

    button_explore.pack()

    # chapter_names = headers.getChapters()
    # createChapterNames()

    # print(chapter_names)

    # Button(root, text="Generate Document", command=generateXML).pack(pady=20)
    Button(root, text='Generate Document', command=generateXML, bg="green", fg="white",
           font=('calibre', 12, 'bold'), padx=10, pady=5).pack(pady=20)

    container.pack(anchor="center")
    canvas.pack(side="left", fill="both", expand=True, anchor="center")
    right.pack(side="right", fill="both", expand=True, anchor="center")
    scrollbar.pack(side="right", fill="both")

    root.mainloop()

    clearFiles()
