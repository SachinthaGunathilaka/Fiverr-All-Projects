if __name__ == '__main__':
    from docxtpl import DocxTemplate, Listing, RichText
    from docx import Document
    from docxcompose.composer import Composer

    # document = Document('./Header.docx')

    chapter_list = ["l1", "l2", "l3", "l4"]
    table_of_content_list = [['@.', 'Les atouts de notre offre hello sachintha'], ['A.', ''],
                             ['A.0.', 'Les moyens humains dédiés au marché'],
                             ['A.0.1.', 'Le Curriculum Vitae des principaux encadrants'],
                             ['A.0.2.', 'Le personnel de chantier'], ['A.1.', 'Les moyens humains dédiés au marché'],
                             ['A.1.1.', 'Le Curriculum Vitae des principaux encadrants'],
                             ['A.1.2.', 'Le personnel de chantier'], ['A.2.', 'Les moyens humains dédiés au marché'],
                             ['A.2.1.', 'Le Curriculum Vitae des principaux encadrants'],
                             ['A.2.2.', 'Le personnel de chantier'], ['B.', ''],
                             ['B.0.', 'Moyens matériels affectés au chantier'], ['B.0.1.', 'Ressources matérielles'],
                             ['B.0.2.', 'Parc matériel'], ['C.', ''], ['C.0.', 'Les atouts de notre offre'],
                             ['C.1.', ''],
                             ['C.2.', 'Fiches méthodes'], ['C.2.1.', 'Mise en œuvre de béton désactivé'],
                             ['C.2.2.', 'Mise en œuvre d’enrobes a la main'],
                             ['C.2.3.', 'Mise en œuvre d’enrobes au finisseur'],
                             ['C.2.4.', 'Mise en œuvre et réglage de GNT'], ['C.2.5.', 'POSE DE BORDURES'],
                             ['C.2.6.', 'Terrassement en déblais'], ['C.2.7.', 'Réseau pluvial'], ['D.', ''],
                             ['D.0.', 'Atouts de l’offre'], ['D.1.', 'Risques sur chantier'],
                             ['D.1.1.', 'Durant tout le chantier'], ['D.2.', 'La sécurité'],
                             ['D.2.1.',
                              'Organisation signalisation temporaire sous circulation (horizontale et verticale)'],
                             ['D.2.2.', 'Moyens à disposition et vérifications 24h/24h'],
                             ['D.2.3.', 'Homme trafic mis à disposition'],
                             ['D.2.4.', 'Formations et actions spécifiques'],
                             ['D.2.5.', 'Signalisation temporaire'], ['D.2.6.', 'Intervention du laboratoire'],
                             ['D.2.7.', 'Terrassement'], ['D.2.8.', 'Traitement de sol'],
                             ['D.2.9.', "Couche de forme et d'assise"], ['D.2.10.', 'Assainissement'],
                             ['D.2.11.', "Adduction d'eau"], ['D.2.12.', 'Tranchées / Réseaux divers'],
                             ['D.2.13.', "Pose d'éléments préfabriqués"],
                             ['D.2.14.', 'Pose de bordures, pavés, dalles'],
                             ['D.2.15.',
                              'Enduits superficiels non marqués CE (surface < 500m2 ou formulés par le client)'],
                             ['D.2.16.', 'Dallage béton'], ['D.2.17.', 'Rabotage de voirie'], ['D.2.18.', 'Enrobés'],
                             ['D.2.19.', 'Béton extrudé par coffrage glissant (en cours)'], ['E.', ''],
                             ['E.0.', 'Les atouts de notre offre'], ['E.1.', 'Environnement'],
                             ['E.1.1.', 'Réduction des émissions sonores'],
                             ['E.1.2.', 'Qualité de l’air et réduction des émissions de poussières'],
                             ['E.2.', 'Risques et moyens de prévention par tâche'],
                             ['E.2.1.', 'Durant tout le chantier'],
                             ['E.2.2.', 'Intervention du laboratoire'], ['E.2.3.', 'Terrassement'],
                             ['E.2.4.', 'Traitement de sol'], ['E.2.5.', "Couche de forme et d'assise"],
                             ['E.2.6.', 'Assainissement'], ['E.2.7.', "Adduction d'eau"],
                             ['E.2.8.', 'Tranchées / Réseaux divers'], ['E.2.9.', 'Pose de bordures, pavés, dalles'],
                             ['E.2.10.',
                              'Enduits superficiels non marqués CE (surface < 500m2 ou formulés par le client)'],
                             ['E.2.11.', 'Dallage béton'], ['E.2.12.', 'Rabotage de voirie'], ['E.2.13.', 'Enrobés'],
                             ['E.2.14.', 'Béton extrudé par coffrage glissant (en cours)'], ['F.', ''],
                             ['F.0.', 'Les atouts de notre offre'], ['F.1.', 'Gestion des déchets'],
                             ['F.1.1.', 'Gestion des déchets'], ['G.', '']]
    main_title = table_of_content_list[0][1]
    print(main_title)

    table_str = RichText()

    pass_A = False
    for [header_num, title] in table_of_content_list[1:]:
        levels = len(header_num.split(".")) - 2
        if levels == 0:

            if pass_A:
                table_str.add("\n", size=10)

            table_str.add(
                header_num + "    " + "." * (108) + "\n",
                color="003399", size=32, italic=True, )

            pass_A = True

        elif levels == 1:
            table_str.add("\n", size=10)
            table_str.add("    " * levels + header_num + "\t" + title + "\n")

        else:
            table_str.add("    " * levels + header_num + "\t" + title + "\n")
            # table_str += "    " * levels + header_num + "\t" + title + "\n"

    # print(table_str)
    chapter_list_str = chapter_list[0]
    for li in chapter_list[1:]:
        chapter_list_str += "\a" + li

    table_of_contents = [[""]]
    doc = DocxTemplate("Header.docx")

    #
    #

    table_title = RichText(main_title + "\n", color="003399", size=32, )

    context = {'title': main_title, "table_title": table_title, "desc_1": "This is the new description",
               "chapter_list": Listing(chapter_list_str),
               "table_of_content": table_str}
    doc.render(context)
    doc.save("generated_doc.docx")

    # for i in range(len(document.paragraphs)):
    #     if document.paragraphs[i].text == "[Main Title]":
    #         # document.paragraphs[i].add_break()
    #         document.paragraphs[i].text = """[Main Title]"""

    # do.save("new.docx")

    # print(document.sections[0].header.paragraphs[0].text)

    files = ["./generated_doc.docx", "./sample.docx"]
    composed = "./merged.docx"

    result = Document(files[0])
    # result.add_page_break()
    composer = Composer(result)

    for i in range(1, len(files)):
        doc = Document(files[i])

        if i != len(files) - 1:
            doc.add_page_break()

        composer.append(doc)

    composer.save(composed)
