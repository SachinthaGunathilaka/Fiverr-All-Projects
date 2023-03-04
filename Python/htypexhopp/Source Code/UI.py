# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'untitled.ui'
#
# Created by: PyQt5 UI code generator 5.15.4
#
# WARNING: Any manual changes made to this file will be lost when pyuic5 is
# run again.  Do not edit this file unless you know what you are doing.


from PyQt5 import QtCore, QtGui, QtWidgets, sip
import docx

import os
import docx
import re
from tkinter import *
from tkinter import ttk

from PyQt5.QtCore import QRect, QSize
from PyQt5.QtGui import QGuiApplication
from PyQt5.QtWidgets import *
from PyQt5 import QtCore
from PyQt5 import QtGui
from PyQt5.QtWidgets import QFileDialog
from bs4 import BeautifulSoup
import shutil
import zipfile
from tkinter import filedialog
from tkinter.filedialog import asksaveasfile
from docxtpl import DocxTemplate, Listing, RichText
from docx import Document
from docxcompose.composer import Composer
from functools import partial

from Header import headers, iter_heading, addHeaderNumbering, initDocFile, getParaIds, zip_file_path, \
    output_doc_path, xml_file_path, output_zip_path, clearFiles


def closeEvent(event):
    clearFiles()


class Ui_Form(object):
    def setupUi(self, Form):
        self.select_all_status = False
        self.check_boxes = []
        self.chapter_titles = []
        Form.setObjectName("Form")
        Form.resize(1600, 900)
        # Form.setMinimumSize(QtCore.QSize(1920, 1080))
        # Form.setMaximumSize(QtCore.QSize(1920, 1080))
        Form.setStyleSheet("background-color: rgb(193, 193, 193);")
        self.frame = QtWidgets.QFrame(Form)
        self.frame.setGeometry(QtCore.QRect(0, 10, 1920, 1080))
        self.frame.setStyleSheet("background-color: rgb(159, 159, 159);")
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        self.label = QtWidgets.QLabel(self.frame)
        self.label.setGeometry(QtCore.QRect(120, 20, 301, 161))
        self.label.setStyleSheet("background-image: url(images/book.jpg);")
        self.label.setText("")
        self.label.setObjectName("label")

        # fname = QFileDialog.getOpenFileName(self, 'Open file',
        #                                     'c:\\', "Image files (*.jpg *.gif)")

        #

        self.insert_btn = QtWidgets.QPushButton(self.frame)
        self.insert_btn.setGeometry(QtCore.QRect(590, 30, 301, 51))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.insert_btn.setFont(font)
        self.insert_btn.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.insert_btn.setObjectName("insert_btn")
        self.insert_btn.clicked.connect(self.insertDocument)

        self.generate_doc_btn = QtWidgets.QPushButton(self.frame)
        self.generate_doc_btn.setGeometry(QtCore.QRect(590, 110, 301, 51))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.generate_doc_btn.setFont(font)
        self.generate_doc_btn.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.generate_doc_btn.setObjectName("generate_doc_btn")
        self.generate_doc_btn.clicked.connect(self.generateXML)

        self.label_2 = QtWidgets.QLabel(self.frame)
        self.label_2.setGeometry(QtCore.QRect(1060, 70, 171, 41))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.label_2.setFont(font)
        self.label_2.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.label_2.setAlignment(QtCore.Qt.AlignCenter)
        self.label_2.setObjectName("label_2")

        self.lineEdit_title = QtWidgets.QLineEdit(self.frame)
        self.lineEdit_title.setGeometry(QtCore.QRect(1240, 70, 571, 41))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.lineEdit_title.setFont(font)
        self.lineEdit_title.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.lineEdit_title.setText("")
        self.lineEdit_title.setAlignment(QtCore.Qt.AlignCenter)
        self.lineEdit_title.setObjectName("lineEdit_title")

        # ----------------------------------------------------------------------
        # Selected Chapter List
        rec = QGuiApplication.screens()[0].geometry()
        scrollarea_width = rec.width() // 2 - 30
        scrollarea_height = rec.height() - 320

        self.scrollArea = QtWidgets.QScrollArea(self.frame)
        self.scrollArea.setGeometry(QtCore.QRect(20, 210, scrollarea_width, scrollarea_height))
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")



        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 898, 2000))
        self.scrollAreaWidgetContents.setMinimumSize(QtCore.QSize(898, 10000))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")

        self.left_frame = QtWidgets.QFrame(self.scrollAreaWidgetContents)
        self.left_frame.setGeometry(QtCore.QRect(0, 0, 931, 2000))
        self.left_frame.setMinimumSize(QtCore.QSize(931, 10000))
        self.left_frame.setStyleSheet("background-color:rgb(197, 204, 208)")
        self.left_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.left_frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.left_frame.setObjectName("left_frame")

        self.scrollArea.setWidget(self.scrollAreaWidgetContents)

        self.show_selected_btn = QtWidgets.QPushButton(self.left_frame)
        self.show_selected_btn.setGeometry(QtCore.QRect(320, 20, 250, 41))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.show_selected_btn.setFont(font)
        self.show_selected_btn.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.show_selected_btn.setObjectName("show_selected_btn")
        self.show_selected_btn.clicked.connect(self.showSelectedChapters)

        # ----------------------------------------------------------------------

        # ----------------------------------------------------------------------

        # All chapter list
        self.scrollArea_2 = QtWidgets.QScrollArea(self.frame)
        # self.scrollArea_2.setGeometry(QtCore.QRect(940, 210, 951, 851))
        self.scrollArea_2.setGeometry(QtCore.QRect(rec.width() // 2, 210, scrollarea_width, scrollarea_height))
        self.scrollArea_2.setWidgetResizable(True)
        self.scrollArea_2.setObjectName("scrollArea_2")

        self.scrollAreaWidgetContents_2 = QtWidgets.QWidget()
        self.scrollAreaWidgetContents_2.setGeometry(QtCore.QRect(0, 0, 898, 2000))
        self.scrollAreaWidgetContents_2.setMinimumSize(QtCore.QSize(898, 10000))
        # self.scrollAreaWidgetContents_2.setMaximumSize(QtCore.QSize(10000, 20000))
        self.scrollAreaWidgetContents_2.setObjectName("scrollAreaWidgetContents_2")

        self.right_frame = QtWidgets.QFrame(self.scrollAreaWidgetContents_2)
        self.right_frame.setGeometry(QtCore.QRect(0, 0, 951, 2000))
        self.right_frame.setMinimumSize(QtCore.QSize(951, 10000))
        # self.right_frame.setMaximumSize(QtCore.QSize(10000, 20000))
        self.right_frame.setStyleSheet("background-color: rgb(226, 219, 210)")
        self.right_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.right_frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.right_frame.setObjectName("right_frame")
        # ----------------------------------------------------------------------

        # Select/Deselect All Button
        self.select_deselect_all_btn = QtWidgets.QPushButton(self.right_frame)
        self.select_deselect_all_btn.setGeometry(QtCore.QRect(360, 20, 201, 41))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.select_deselect_all_btn.setFont(font)
        self.select_deselect_all_btn.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.select_deselect_all_btn.setObjectName("select_deselect_all_btn")
        self.select_deselect_all_btn.clicked.connect(self.selectAll)

        self.scrollArea_2.setWidget(self.scrollAreaWidgetContents_2)
        self.retranslateUi(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)

        # self.checkBoxEvent()

    def insertDocument(self):
        clearFiles()

        fname = QFileDialog.getOpenFileName(self.insert_btn, 'Open file',
                                            '/', "Doc files (*.docx)")[0]

        if len(fname) == 0:
            return
        self.doc_file_path = fname
        # print(self.doc_file_path)
        initDocFile(self.doc_file_path)

        # self.doc_file_path = "Word Document.docx"
        # initDocFile("Word Document.docx")

        self.showAllChapterNames()

    def showAllChapterNames(self):

        # doc_file_path = "Word Document.docx"
        doc = docx.Document(self.doc_file_path)
        addHeaderNumbering(doc)

        width = 600
        height = 30
        x_pos = 50
        y_pos = 70
        y_increment_val_1 = 50
        y_increment_val_2 = 40

        _translate = QtCore.QCoreApplication.translate

        index = 0
        for header_ in headers.headers:
            level = headers.getLevels(header_[0])

            inc_val = y_increment_val_1
            if level > 0:
                inc_val = y_increment_val_2

            y_pos += inc_val

            checkbox = QtWidgets.QCheckBox(self.right_frame)
            checkbox.setGeometry(QtCore.QRect(x_pos * (level + 1), y_pos, width, height))
            font = QtGui.QFont()
            font.setFamily("Times New Roman")
            font.setPointSize(12)
            checkbox.setFont(font)
            checkbox.setStyleSheet("")
            checkbox.setObjectName("checkBox")

            checkbox.setText(_translate("Form", header_[0] + " " + header_[1]))
            checkbox.stateChanged.connect(partial(self.checkBoxEvent, index))

            checkbox.show()
            self.check_boxes.append(checkbox)

            index += 1

    # def checkBoxEvent(self, select_all_mode):
    #     if not select_all_mode:

    def selectAll(self):
        self.select_all_status = not self.select_all_status
        for checkbox in self.check_boxes:
            checkbox.setChecked(self.select_all_status)

    def checkBoxEvent(self, index):
        level = headers.getLevels(headers.headers[index][0])
        if level == 0:
            return
        for i in range(index + 1, len(self.check_boxes)):
            if headers.getLevels(headers.headers[i][0]) <= level:
                break

            self.check_boxes[i].setChecked(self.check_boxes[index].isChecked())

        # if self.check_boxes[index].isChecked():
        #     parent_boxes = 0
        #     for i in range(index-1, -1, -1):
        #         if parent_boxes == level:
        #             break
        #
        #         if headers.getLevels(headers.headers[i][0]) == level-parent_boxes-1:
        #             self.check_boxes[i].setChecked(True)
        #             parent_boxes += 1

    def showSelectedChapters(self):

        # chapter_order contains the indices of the selected chapters (In order)
        # We can locate to a chapter by headers.headers[chapter_order_element]

        self.chapter_order = []
        self.chapter_titles = []
        self.chapter_titles_line_edit = []

        for child in self.left_frame.children():
            child.deleteLater()

        _translate = QtCore.QCoreApplication.translate

        self.show_selected_btn = QtWidgets.QPushButton(self.left_frame)
        self.show_selected_btn.setGeometry(QtCore.QRect(320, 20, 250, 41))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.show_selected_btn.setFont(font)
        self.show_selected_btn.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.show_selected_btn.setObjectName("show_selected_btn")
        self.show_selected_btn.clicked.connect(self.showSelectedChapters)
        self.show_selected_btn.show()
        self.show_selected_btn.setText(_translate("Form", "Afficher"))

        # -----------------------------------------------------------
        # Selected Chapter Details
        # -----------------------------------------------------------

        index = 0
        for i in range(len(headers.headers)):
            if self.check_boxes[i].isChecked():
                self.chapter_order.append(i)
                x_pos = 70 + 60 * headers.getLevels(headers.headers[i][0])
                y_pos = 110 + 60 * index
                line_edit = QtWidgets.QLineEdit(self.left_frame)
                line_edit.setGeometry(QtCore.QRect(x_pos, y_pos, 571, 41))
                font = QtGui.QFont()
                font.setFamily("Times New Roman")
                font.setPointSize(12)
                font.setBold(True)
                font.setWeight(75)
                line_edit.setFont(font)
                line_edit.setStyleSheet("background-color: rgb(255, 255, 255);")
                line_edit.setAlignment(QtCore.Qt.AlignCenter)
                line_edit.setObjectName("lineEdit_2")

                line_edit.setText(_translate("Form", headers.headers[i][0] + " " + headers.headers[i][1]))

                up_btn = QtWidgets.QPushButton(self.left_frame)
                up_btn.setGeometry(QtCore.QRect(x_pos + 590, y_pos, 31, 21))
                up_btn.setStyleSheet("background-color: rgb(255, 255, 255);\n"
                                     "image: url(images/up-arrow.png);")
                up_btn.setText("")
                up_btn.setObjectName("pushButton_4")
                up_btn.clicked.connect(partial(self.goUp, i))

                down_btn = QtWidgets.QPushButton(self.left_frame)
                down_btn.setGeometry(QtCore.QRect(x_pos + 590, y_pos + 20, 31, 21))
                down_btn.setStyleSheet("background-color: rgb(255, 255, 255);\n"
                                       "image: url(images/down-arrow.png);")
                down_btn.setText("")
                down_btn.setObjectName("pushButton_5")
                down_btn.clicked.connect(partial(self.goDown, i))

                line_edit.show()
                up_btn.show()
                down_btn.show()

                self.chapter_titles_line_edit.append(line_edit)

                index += 1

        # -----------------------------------------------------------------------

    def rebuildSelectedChapters(self):
        self.chapter_titles = []
        for line_edit in self.chapter_titles_line_edit:
            # print(line_edit.text())
            self.chapter_titles.append(line_edit.text())
        for child in self.left_frame.children():
            child.deleteLater()

        _translate = QtCore.QCoreApplication.translate

        self.show_selected_btn = QtWidgets.QPushButton(self.left_frame)
        self.show_selected_btn.setGeometry(QtCore.QRect(320, 20, 250, 41))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.show_selected_btn.setFont(font)
        self.show_selected_btn.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.show_selected_btn.setObjectName("show_selected_btn")
        self.show_selected_btn.clicked.connect(self.showSelectedChapters)
        self.show_selected_btn.show()
        self.show_selected_btn.setText(_translate("Form", "Afficher"))

        # -----------------------------------------------------------
        # Selected Chapter Details
        # -----------------------------------------------------------

        index = 0
        for i in self.chapter_order:
            x_pos = 70 + 60 * headers.getLevels(headers.headers[i][0])
            y_pos = 110 + 60 * index
            line_edit = QtWidgets.QLineEdit(self.left_frame)
            line_edit.setGeometry(QtCore.QRect(x_pos, y_pos, 571, 41))
            font = QtGui.QFont()
            font.setFamily("Times New Roman")
            font.setPointSize(12)
            font.setBold(True)
            font.setWeight(75)
            line_edit.setFont(font)
            line_edit.setStyleSheet("background-color: rgb(255, 255, 255);")
            line_edit.setAlignment(QtCore.Qt.AlignCenter)
            line_edit.setObjectName("lineEdit_2")

            # line_edit.setText(_translate("Form", self.chapter_titles[index]))
            line_edit.setText(_translate("Form", self.chapter_titles[index]))

            up_btn = QtWidgets.QPushButton(self.left_frame)
            up_btn.setGeometry(QtCore.QRect(x_pos + 590, y_pos, 31, 21))
            up_btn.setStyleSheet("background-color: rgb(255, 255, 255);\n"
                                 "image: url(images/up-arrow.png);")
            up_btn.setText("")
            up_btn.setObjectName("pushButton_4")
            up_btn.clicked.connect(partial(self.goUp, i))

            down_btn = QtWidgets.QPushButton(self.left_frame)
            down_btn.setGeometry(QtCore.QRect(x_pos + 590, y_pos + 20, 31, 21))
            down_btn.setStyleSheet("background-color: rgb(255, 255, 255);\n"
                                   "image: url(images/down-arrow.png);")
            down_btn.setText("")
            down_btn.setObjectName("pushButton_5")
            down_btn.clicked.connect(partial(self.goDown, i))

            line_edit.show()
            up_btn.show()
            down_btn.show()

            self.chapter_titles_line_edit[index] = line_edit
            index += 1

    def goUp(self, index):
        i = 0
        while i < len(self.chapter_order):
            if self.chapter_order[i] == index:
                break
            i += 1

        if i == 0:
            return

        before_chap_start = i - 1
        while before_chap_start != -1:

            if headers.getLevels(headers.headers[self.chapter_order[i]][0]) >= headers.getLevels(
                    headers.headers[self.chapter_order[before_chap_start]][0]):
                break

            before_chap_start -= 1

        before_chap_end = before_chap_start + 1
        while before_chap_end != len(self.chapter_order):
            if headers.getLevels(headers.headers[self.chapter_order[i]][0]) >= headers.getLevels(
                    headers.headers[self.chapter_order[before_chap_end]][0]):
                break

            before_chap_end += 1

        after_index = i + 1
        while after_index != len(self.chapter_order):
            if headers.getLevels(headers.headers[self.chapter_order[i]][0]) >= headers.getLevels(
                    headers.headers[self.chapter_order[after_index]][0]):
                break

            after_index += 1

        # print(before_chap_start, before_chap_end, i, after_index)

        chaps = self.chapter_order[i:after_index]
        chaps.extend(self.chapter_order[before_chap_end:i])
        chaps.extend(self.chapter_order[before_chap_start:before_chap_end])

        self.chapter_order[before_chap_start:after_index] = chaps

        chaps_title = self.chapter_titles_line_edit[i:after_index]
        chaps_title.extend(self.chapter_titles_line_edit[before_chap_end:i])
        chaps_title.extend(self.chapter_titles_line_edit[before_chap_start:before_chap_end])

        self.chapter_titles_line_edit[before_chap_start:after_index] = chaps_title
        # print(chaps)
        self.rebuildSelectedChapters()

    def goDown(self, index):
        i = 0
        while i < len(self.chapter_order):
            if self.chapter_order[i] == index:
                break
            i += 1

        if i == len(self.chapter_order) - 1:
            return

        after_chap_start = i + 1
        while after_chap_start != len(self.chapter_order):

            if headers.getLevels(headers.headers[self.chapter_order[i]][0]) >= headers.getLevels(
                    headers.headers[self.chapter_order[after_chap_start]][0]):
                break

            after_chap_start += 1

        after_chap_end = after_chap_start + 1
        while after_chap_end != len(self.chapter_order):
            if headers.getLevels(headers.headers[self.chapter_order[i]][0]) >= headers.getLevels(
                    headers.headers[self.chapter_order[after_chap_end]][0]):
                break

            after_chap_end += 1

        after_index = i + 1
        while after_index != len(self.chapter_order):
            if headers.getLevels(headers.headers[self.chapter_order[i]][0]) >= headers.getLevels(
                    headers.headers[self.chapter_order[after_index]][0]):
                break

            after_index += 1

        # print(after_chap_start, after_chap_end, i, after_index)

        chaps = self.chapter_order[after_chap_start:after_chap_end]
        chaps.extend(self.chapter_order[after_index:after_chap_start])
        chaps.extend(self.chapter_order[i:after_index])

        self.chapter_order[i:after_chap_end] = chaps

        chaps_title = self.chapter_titles_line_edit[after_chap_start:after_chap_end]
        chaps_title.extend(self.chapter_titles_line_edit[after_index:after_chap_start])
        chaps_title.extend(self.chapter_titles_line_edit[i:after_index])

        self.chapter_titles_line_edit[i:after_chap_end] = chaps_title
        # print(chaps)
        self.rebuildSelectedChapters()

    def generateXML(self):

        try:
            self.chapter_titles = []
            for line_edit in self.chapter_titles_line_edit:
                # print(line_edit.text())
                self.chapter_titles.append(line_edit.text())

            try:
                shutil.rmtree("Word Document")

            except:
                pass

            with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
                zip_ref.extractall("Word Document")

            for file in [output_doc_path, output_zip_path, "pre_output.docx"]:
                if os.path.exists(file):
                    os.remove(file)

            if os.path.exists(xml_file_path):
                selectedData = []

                for chapter_index in self.chapter_order:
                    selectedData.append([chapter_index, headers.getHeaderPath(chapter_index)])

                header_para_ids = getParaIds()

                xml_string = open(xml_file_path).read()
                body_start = re.search("<w:body[\s\S]*>", xml_string).start()

                body_end = body_start + 5
                while True:
                    if xml_string[body_end] == '<':
                        break
                    body_end += 1

                out_content = xml_string[:body_end]

                # out_content += open("empty_template.txt").read()
                for input_header_details in selectedData:

                    start_paraId = header_para_ids[input_header_details[0]]

                    end_paraId = None
                    if input_header_details[0] != len(headers.headers) - 1:
                        end_paraId = header_para_ids[input_header_details[0] + 1]

                    header_start = re.search(f"paraId=\"{start_paraId}\"", xml_string).start()

                    header_1_start = header_start
                    while True:
                        if xml_string[header_1_start] == "<":
                            break

                        header_1_start -= 1

                    if input_header_details[0] != len(headers.headers) - 1:
                        header_end = re.search(f"paraId=\"{end_paraId}\"", xml_string).start()
                        header_2_start = header_end

                        while True:
                            if xml_string[header_2_start] == "<":
                                break

                            header_2_start -= 1

                        out_content += xml_string[header_1_start:header_2_start]

                    else:

                        last_pos = re.search("</w:body[\s\S]*>", xml_string).start()
                        out_content += xml_string[header_1_start:last_pos]

                end_content = re.search("</w:body[\s\S]*>", xml_string).start()
                out_content += xml_string[end_content:]

                out_xml = open(xml_file_path, 'w')

                out_xml.write(out_content)
                out_xml.close()

                self.CreateDoc()


        except:
            print("Some error occured")

    def CreateDoc(self):
        shutil.make_archive("output", 'zip', "Word Document")
        os.rename('output.zip', 'pre_output.docx')

        self.changeTitleNames()

        final_doc_fila_path = QFileDialog.getSaveFileName(self.insert_btn, 'Save Document As',
                                                          '/', "Doc files (*.docx)")[0]

        if len(final_doc_fila_path) == 0:
            return

        self.createFinalDocFromTemplate(output_doc_path, final_doc_fila_path)

    def changeTitleNames(self):
        document = Document("pre_output.docx")

        index = 0
        for i in range(len(document.paragraphs)):
            isItHeading = re.match('Heading ([1])|Heading ([2])|Heading ([3])', document.paragraphs[i].style.name)
            if isItHeading is not None:
                # print(self.chapter_titles[index])

                # print(int(isItHeading.groups()[0]), document.paragraphs[i].text.strip())
                document.paragraphs[i].text = self.chapter_titles[index]
                # print(self.chapter_titles[index])
                index += 1

        document.save("output.docx")

    def createFinalDocFromTemplate(self, content_doc, output_doc):

        main_title = self.lineEdit_title.text()

        first_chapter = False
        chapter_list_str = ""

        index = 0
        for order in self.chapter_order:
            if headers.getLevels(headers.headers[order][0]) == 0:

                if not first_chapter:
                    first_chapter = True
                    chapter_list_str = self.chapter_titles[index]

                else:
                    chapter_list_str += "\a" + self.chapter_titles[index]

            index += 1

        doc = DocxTemplate("Header.docx")

        table_title = RichText(main_title + "\n", color="003399", size=32, )

        context = {'title': main_title, "table_title": table_title,
                   "chapter_list": Listing(chapter_list_str)}
        doc.render(context)
        doc.save("generated_doc.docx")

        files = ["./generated_doc.docx", content_doc]
        result = Document(files[0])
        result.add_page_break()

        composer = Composer(result)

        for i in range(1, len(files)):
            doc = Document(files[i])

            if i != len(files) - 1:
                doc.add_page_break()

            composer.append(doc)

        composer.save(output_doc)

    def retranslateUi(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Mémoire technique"))
        self.insert_btn.setText(_translate("Form", "Inserer le document"))

        self.generate_doc_btn.setText(_translate("Form", "Générer le document"))
        self.label_2.setText(_translate("Form", "Titre "))

        self.select_deselect_all_btn.setText(_translate("Form", "cocher/décocher"))
        self.show_selected_btn.setText(_translate("Form", "Afficher"))


if __name__ == "__main__":
    import sys

    clearFiles()

    app = QtWidgets.QApplication(sys.argv)
    app.setWindowIcon(QtGui.QIcon("images/icon.ico"))

    Form = QtWidgets.QWidget()
    ui = Ui_Form()
    ui.setupUi(Form)
    Form.show()

    app.exec_()
    clearFiles()
    sys.exit()
